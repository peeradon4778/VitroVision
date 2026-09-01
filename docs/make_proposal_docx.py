# -*- coding: utf-8 -*-
"""แปลง docs/proposal_th_draft.md → docx สไตล์ YSC Proposal
- ฟอนต์ TH SarabunPSK ทั้งเล่ม · หัวข้อ bold 16pt · เนื้อหา 16pt ind_left
- ตาราง style TableGrid · กระดาษ A4 ขอบ 1" · footer หมายเลขหน้า

รัน:  python make_proposal_docx.py [--in proposal_th_draft.md] [--out proposal_th_draft.docx]
"""
import argparse
import pathlib
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = 'TH SarabunPSK'
HERE = pathlib.Path(__file__).resolve().parent
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def set_font(run, bold=False, italic=False, size=16):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(a), FONT)


def add_rich(p, text, size=16, base_bold=False):
    """เพิ่ม run โดยแยก **bold** และ *italic* จากข้อความ"""
    for part in re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = p.add_run(part[2:-2])
            set_font(r, bold=True, italic=base_bold, size=size)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_font(r, bold=base_bold, italic=True, size=size)
        else:
            r = p.add_run(part)
            set_font(r, bold=base_bold, size=size)


def set_para(p, before=0, after=120, ind_left=None, hanging=None, jc=None):
    pf = p.paragraph_format
    if before:
        pf.space_before = Pt(before / 20) if False else None
        # ใช้ twips ผ่าน XML ตรง ๆ (Word ไทยนิยม twips)
        pPr = p._p.get_or_add_pPr()
        sp = pPr.find(W + 'spacing')
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(W + 'before', str(before))
        sp.set(W + 'after', str(after))
    else:
        pPr = p._p.get_or_add_pPr()
        sp = pPr.find(W + 'spacing')
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(W + 'after', str(after))
    if ind_left is not None:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(W + 'ind')
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(W + 'left', str(ind_left))
        if hanging:
            ind.set(W + 'hanging', str(hanging))
    if jc:
        p.alignment = jc


def add_footer_page(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run('หน้า ')
    set_font(r1, size=14)
    fld = OxmlElement('w:fldSimple')
    fld.set(W + 'instr', 'PAGE')
    p._p.append(fld)
    r2 = p.add_run()
    set_font(r2, size=14)


def build_table(doc, rows):
    """rows: list[list[str]] — แถวแรกเป็น header"""
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = tbl.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            txt = row[j] if j < len(row) else ''
            add_rich(p, txt, size=14, base_bold=(i == 0))
    # ตั้งความกว้างอัตโนมัติ
    tbl.autofit = True
    return tbl


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--in', dest='src', default=HERE / 'proposal_th_draft.md')
    ap.add_argument('--out', dest='out', default=HERE / 'proposal_th_draft.docx')
    args = ap.parse_args()

    src = pathlib.Path(args.src)
    text = src.read_text(encoding='utf-8')

    doc = Document()
    # กระดาษ A4 + ขอบ 1 นิ้ว
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    add_footer_page(doc)

    lines = text.splitlines()
    i = 0
    in_table = False
    table_buf = []
    while i < len(lines):
        line = lines[i].rstrip()

        # ตาราง: | ... |
        if line.strip().startswith('|') and line.strip().endswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if set(''.join(cells).replace('-', '').replace(':', '').replace(' ', '')) == set() and all('-' in c for c in cells):
                i += 1  # ข้าม separator row
                continue
            table_buf.append(cells)
            in_table = True
            i += 1
            continue
        if in_table:
            build_table(doc, table_buf)
            table_buf = []
            in_table = False

        # ขีดคั่น
        if re.match(r'^-{3,}$', line.strip()):
            i += 1
            continue
        # หมายเหตุ/blockquote
        if line.startswith('>'):
            p = doc.add_paragraph()
            add_rich(p, line.lstrip('>').strip(), size=14)
            set_para(p, after=120, ind_left=432)
            i += 1
            continue
        # หัวข้อ
        m = re.match(r'^(#{1,3})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            p = doc.add_paragraph()
            add_rich(p, m.group(2), size=16, base_bold=True)
            set_para(p, before=280 if level <= 2 else 160, after=120)
            i += 1
            continue
        # รายการ
        m = re.match(r'^(\d+)\.\s+(.*)$', line)
        if m:
            p = doc.add_paragraph()
            add_rich(p, f'{m.group(1)}. {m.group(2)}', size=16)
            set_para(p, after=80, ind_left=720)
            i += 1
            continue
        m = re.match(r'^[-*]\s+(.*)$', line)
        if m:
            p = doc.add_paragraph()
            add_rich(p, '• ' + m.group(1), size=16)
            set_para(p, after=80, ind_left=720)
            i += 1
            continue
        # รูปภาพ (flowchart)
        m = re.match(r'^!\[[^\]]*\]\(([^)]+)\)\s*$', line)
        if m:
            rel = m.group(1).replace('\\', '/')
            img = (HERE.parent / rel) if not rel.startswith(('/', 'C:')) else pathlib.Path(rel)
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(str(img), width=Cm(16))
                set_para(p, before=120, after=120)
            else:
                print(f'[warn] ไม่พบรูป {img}')
            i += 1
            continue
        # เนื้อหาว่าง
        if not line.strip():
            i += 1
            continue
        # ย่อหน้าปกติ
        p = doc.add_paragraph()
        add_rich(p, line, size=16)
        # บรรณานุกรม: ขึ้นต้นด้วยชื่อผู้แต่ง+ปี → hanging
        if re.match(r'^[A-Z][A-Za-zÀ-ÿ\'\-]+,', line) or line.startswith('ศูนย์พันธุวิศวกรรม'):
            set_para(p, after=120, ind_left=720, hanging=720)
        else:
            set_para(p, after=120, ind_left=432)
        i += 1

    if table_buf:
        build_table(doc, table_buf)

    out = pathlib.Path(args.out)
    doc.save(str(out))
    print(f'[OK] สร้าง {out} แล้ว')


if __name__ == '__main__':
    main()
