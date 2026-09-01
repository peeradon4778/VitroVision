# -*- coding: utf-8 -*-
"""copy_hf — คัดลอก header/footer (ตาราง+โลโก้) จาก template ทางการ YSC ใส่ proposal_ysc.docx"""
import copy
import re

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part

TEMPLATE = "docs/_ysc_template/YSC-Proposal_Template_200726.docx"
TARGET = "docs/proposal_ysc.docx"


def copy_hf(src_sec, dst_sec, which):
    """which = 'header' หรือ 'footer'"""
    src_part = getattr(src_sec, which).part
    dst_part = getattr(dst_sec, which).part
    src_el = src_part.element
    dst_el = dst_part.element

    # เก็บ children ต้นฉบับ (tables/paragraphs)
    children = [c for c in list(src_el) if c.tag.endswith("}tbl") or c.tag.endswith("}p")]
    # ล้างเนื้อเก่าของปลายทาง
    for c in list(dst_el):
        if c.tag.endswith("}tbl") or c.tag.endswith("}p"):
            dst_el.remove(c)
    # คัดลอกเข้า
    for c in children:
        dst_el.append(copy.deepcopy(c))

    # คัดลอกรูป (image rels) — ให้ rId เดิมชี้หาภาพที่คัดมา
    for rId, rel in list(src_part.rels.items()):
        if rel.reltype != RT.IMAGE:
            continue
        blob = rel.target_part.blob
        ct = rel.target_part.content_type
        # ชื่อไฟล์ส่วนใหม่
        ext = ".png" if "png" in ct else ".jpg"
        safe = re.sub(r"\D", "", rId)
        partname = PackURI(f"/word/media/hf_copy_{safe}{ext}")
        new_part = Part(partname, ct, blob, dst_part.package)
        dst_part.rels.add_relationship(RT.IMAGE, new_part, rId=rId)


doc_t = Document(TEMPLATE)
doc_o = Document(TARGET)

src_sec = doc_t.sections[0]
dst_sec = doc_o.sections[0]

copy_hf(src_sec, dst_sec, "header")
copy_hf(src_sec, dst_sec, "footer")

doc_o.save(TARGET)
print("[OK] header/footer คัดจาก template ทางการ ->", TARGET)

# ตรวจ
d = Document(TARGET)
s = d.sections[0]
print("header tables:", len(s.header.tables), "| header text:",
      [c.text.strip()[:40] for r in s.header.tables[0].rows for c in r.cells] if s.header.tables else "?")
print("footer tables:", len(s.footer.tables), "| footer text:",
      [c.text.strip()[:40] for r in s.footer.tables[0].rows for c in r.cells] if s.footer.tables else "?")
