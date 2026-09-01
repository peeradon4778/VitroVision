# -*- coding: utf-8 -*-
"""fix_footer — ลบ footer เก่า (Version 3...) แล้วใส่เลขหน้ากลาง (PAGE field) เท่านั้น"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

TARGET = "docs/proposal_ysc.docx"
d = Document(TARGET)
footer = d.sections[0].footer
fel = footer._element

# ลบตาราง + ย่อหน้าเก่าทั้งหมดใน footer
for t in list(fel.findall(qn("w:tbl"))):
    fel.remove(t)
for p in list(fel.findall(qn("w:p"))):
    fel.remove(p)

# สร้างย่อหน้าเลขหน้า ตรงกลาง
p = footer.add_paragraph()
p.alignment = 1  # CENTER


def _run(text=None, font="TH Sarabun New", size=14):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rF = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(a), font)
    rPr.append(rF)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
    r.append(rPr)
    if text is not None:
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
    return r


r1 = _run(); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1.append(f1)
r2 = _run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "; r2.append(it)
r3 = _run(); f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end"); r3.append(f2)
r4 = _run("1")  # ค่าเริ่มต้นให้ Word โชว์เลข

p._p.append(r1)
p._p.append(r2)
p._p.append(r3)
p._p.append(r4)

d.save(TARGET)
print("[OK] footer -> เลขหน้ากลางเท่านั้น (ไม่มี Version) :", TARGET)

d2 = Document(TARGET)
ftxt = [p.text for p in d2.sections[0].footer.paragraphs if p.text.strip()]
print("footer text now:", ftxt)
