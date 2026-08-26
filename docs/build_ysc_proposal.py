# -*- coding: utf-8 -*-
"""เติมเนื้อหา VitroVision ลง template YSC Proposal อย่างเป็นทางการ (แก้ template โดยตรง)
ผล = .docx ที่แก้ไขได้ใน Word ทุกตัวอักษร · ฟอนต์ TH Sarabun PSK 16pt
รัน: python build_ysc_proposal.py
"""
import os
import docx
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
IN = os.path.join(HERE, "_ysc_template", "YSC-Proposal_Template_200726.docx")
OUT = os.path.join(HERE, "ysc_proposal_filled.docx")
FIG = os.path.join(HERE, "assets", "vitro_architecture.png")
FONT = "TH Sarabun PSK"
BODY = 16
J = {None: None, "c": 1, "l": 0, "r": 2, "j": 3}

doc = docx.Document(IN)
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# ---------------------------------------------------------------- helpers
def set_font(run, bold=False, size=BODY, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), FONT)
    for a in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rFonts.get(qn(a)) is not None:
            del rFonts.attrib[qn(a)]


def make_para(text, bold=False, size=BODY, italic=False, jc=None, ind_left=None, after=120):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, bold=bold, size=size, italic=italic)
    pf = p.paragraph_format
    pf.space_after = Pt(after / 20)
    if jc:
        pf.alignment = J[jc]
    if ind_left is not None:
        pf.left_indent = Cm(ind_left)
    return p


def move_after(anchor_el, obj_el):
    anchor_el.addnext(obj_el)
    return obj_el


_GRAY = {"7F7F7F", "7f7f7f"}


def is_gray(pp):
    for r in pp.runs:
        try:
            if r.font.color and r.font.color.type is not None and str(r.font.color.rgb) in _GRAY:
                return True
        except Exception:
            pass
    return False


def is_placeholder(pp):
    t = pp.text.strip()
    if not t:
        return False
    if is_gray(pp):
        return True
    if t.startswith("(ใส่ชื่อหัวข้อ)") or t.startswith("(กรุณาลบข้อความสีเทา"):
        return True
    if t.startswith("รูปที่") and "XXX" in t or t.startswith("ตารางที่") and "XXX" in t:
        return True
    return False


def clear_section_after(heading_para):
    el = heading_para._p
    nxt = el.getnext()
    while nxt is not None:
        if nxt.tag == qn("w:p"):
            pp = Paragraph(nxt, heading_para._parent)
            if pp.style.name == "List Paragraph" and pp.text.strip() and not is_placeholder(pp):
                break
            if is_placeholder(pp):
                to_del = nxt; nxt = nxt.getnext(); to_del.getparent().remove(to_del); continue
        nxt = nxt.getnext()


def add_blocks(anchor_para, blocks):
    anchor = anchor_para._p
    for blk in blocks:
        kind = blk[0]
        if kind == "p":
            _, text, opts = blk[0], blk[1], (blk[2] if len(blk) > 2 else {})
            p = make_para(text, **opts)
            anchor = move_after(anchor, p._p)
        elif kind == "fig":
            p_img = doc.add_paragraph()
            p_img.add_run().add_picture(FIG, width=Cm(15))
            p_img.alignment = 1
            anchor = move_after(anchor, p_img._p)
            p_cap = make_para(blk[1], bold=True, size=12, jc="c", after=60)
            anchor = move_after(anchor, p_cap._p)
        elif kind == "note":
            p = make_para(blk[1], italic=True, size=12)
            anchor = move_after(anchor, p._p)
    return anchor


# ---------------------------------------------------------------- หัวข้อ
SECTION_HEADINGS = [
    "หลักการและเหตุผล (Rationale)",
    "วัตถุประสงค์ (Objective/s) / เป้าหมายทางวิศวกรรม (Engineering Goal/s)",
    "สมมติฐาน (Hypothesis/es)",
    "ประโยชน์และผลที่คาดว่าจะได้รับ (Benefits and Expected Results)",
    "วัสดุอุปกรณ์และสถานที่ดำเนินงาน (Materials and Workplace/s)",
    "ระเบียบวิธีการทดลอง (Methodology)",
    "การวิเคราะห์ข้อมูล (Data Analysis)",
    "ความเสี่ยงและความปลอดภัย (Risk and Safety)",
    "แผนการดำเนินงาน (Research Plan)",
    "บรรณานุกรม (Bibliography)",
]

# ---- เนื้อหาแต่ละ section (real จาก proposal_th_draft) ----
C = {
"หลักการและเหตุผล (Rationale)": [
    ("p", "การเพาะเลี้ยงเนื้อเยื่อพืช (plant tissue culture / micropropagation) เป็นเทคโนโลยีการขยายพันธุ์พืชที่ใช้ชิ้นส่วนขนาดเล็กเพาะเลี้ยงในสภาพปลอดเชื้อบนอาหารสังเคราะห์ และถูกใช้ในเชิงพาณิชย์ครอบคลุมพืชเกษตร อาหาร เภสัชกรรม และเครื่องสำอางทั่วโลก (Hasnain et al., 2022; Chandran et al., 2020) แม้เทคโนโลยีจะพัฒนาไปมาก แต่กระบวนการที่ยังพึ่งพาแรงงานคนสูงคือการตัดสินใจว่าเมื่อใดจึงควรย้ายต้นกล้าออกจากขวดไปยังสภาพอนุบาล (acclimatization/hardening) ซึ่งต้องพิจารณารายขวดทุก 3–8 สัปดาห์ ขึ้นกับชนิดพืช (Pastelín Solano et al., 2019; Regni et al., 2025)"),
    ("p", "การตัดสินใจที่ผิดพลาด — ย้ายเร็วเกินไป (ต้นยังไม่สมบูรณ์ ระบบรากไม่ดี) หรือช้าเกินไป (แออัด เสี่ยง hyperhydricity) — อาจทำให้เนื้อเยื่อตาย (necrosis) และลดประสิทธิภาพการขยายพันธุ์ (Abdalla et al., 2022) งานวิจัยก่อนหน้านำ computer vision มาใช้วัดพืชในขวดแบบไม่ทำลาย เช่น ระบบ Phenomenon (multi-sensor + random forest) วัด projected area และ canopy height (Bethge et al., 2023) และ Regni et al. (2025) ใช้ภาพ 3D จากสมาร์ตโฟนวัด canopy และ shoot density แต่ระบบดังกล่าวเป็น hardware เฉพาะราคาสูง หรือยังไม่ใช้ foundation model ที่ทำงาน zero-shot ข้ามชนิดพืช"),
    ("p", "ช่องว่างองค์ความรู้: ยังไม่มีระบบ low-cost ที่ติดตามการเจริญของต้นกล้าในขวดตามเวลา (time-series) แบบ end-to-end ตั้งแต่ถ่ายภาพ → แบ่งส่วน → เก็บ feature → วิเคราะห์อัตราการเจริญ → แจ้งเตือนความพร้อมอนุบาล และยังไม่มีงาน zero-shot foundation model ที่ถูกประเมินเทียบกับวิธีพื้นฐาน (baseline) ด้วยชุดข้อมูลภาพผ่านขวดแก้วโดยตรง รวมทั้งยังไม่มีชุดข้อมูลเปิดของภาพเพาะเลี้ยงเนื้อเยื่อผ่านขวดแก้วสำหรับการวิจัยต่อ จึงเป็นที่มาของโครงงานนี้"),
    ("p", "โครงงานนี้ต่อยอดจากการทดลองนำร่องที่พิสูจน์ว่าโมเดล SAM3 แบบ zero-shot (Segment Anything Model 3) สามารถแบ่งส่วนต้นพืชภายในขวดแก้วได้แม้มีแสงสะท้อน ไอน้ำ และความโค้งของแก้ว ด้วยการใช้พรอมป์ข้อความ (text prompt) และปรับใช้ข้ามชนิดพืชได้โดยไม่ต้องฝึกโมเดลใหม่"),
    ("note", "ทั้งนี้ หากเป็นโครงการต่อเนื่อง: โครงงานนี้ต่อยอดจากงานรุ่นก่อนหน้า (VitroVision v1) อย่างชัดเจน โดยเปลี่ยนจากงานเฉพาะพืชไปเป็นระบบ zero-shot ข้ามชนิด + การติดตามตามเวลา และมีองค์ความรู้/ผลการวิจัยใหม่ในปีปัจจุบัน"),
    ("note", "(กรุณาลบข้อความสีเทานี้ทิ้งก่อนนำไปใช้ในเอกสารจริง)"),
],
"วัตถุประสงค์ (Objective/s) / เป้าหมายทางวิศวกรรม (Engineering Goal/s)": [
    ("p", "1. สร้างชุดข้อมูลภาพถ่ายจริงของพืชเพาะเลี้ยงเนื้อเยื่อผ่านขวดแก้ว (≥ 100 ขวด พร้อม metadata และการกำกับภาพโดยมนุษย์บางส่วน) ซึ่งเป็นชุดข้อมูลใหม่ที่ยังไม่มีในงานวิจัยก่อนหน้า", {"ind_left": 0.75}),
    ("p", "2. พัฒนาระบบติดตามการเจริญตามเวลา (time-series monitoring) ถ่ายภาพขวดซ้ำรายรอบ → แบ่งส่วนด้วย SAM3 → เก็บ feature → วิเคราะห์ growth curve → แจ้งเตือนเมื่อต้นพร้อมอนุบาล", {"ind_left": 0.75}),
    ("p", "3. ประเมินเปรียบเทียบกับวิธีพื้นฐาน (SAM2, YOLO-seg, การแบ่งส่วนเชิงคลาสสิก) ด้วย mIoU, Dice, precision, recall, F1 พร้อมวิเคราะห์ความไวของพรอมป์และเกณฑ์", {"ind_left": 0.75}),
    ("p", "4. ตรวจสอบความถูกต้องของระบบกับค่าอ้างอิงจากผู้ประเมิน (ground truth) และปรับปรุงการตรวจจับขวด ROI ที่เป็นข้อจำกัดของผลนำร่อง", {"ind_left": 0.75}),
],
"สมมติฐาน (Hypothesis/es)": [
    ("p", "H1 (เชิงเทคนิค — segmentation): SAM3 ที่ใช้ text prompts 5 คำ (plant, leaf, shoot, stem, root) สามารถแบ่งส่วนต้นพืชเพาะเลี้ยงเนื้อเยื่อผ่านขวดแก้วที่มี glare, condensation และ reflection ได้ โดยมี mIoU เฉลี่ย ≥ 0.65 เทียบกับ ground truth ที่ annotate โดยมนุษย์ และสูงกว่า baseline (SAM2, YOLO-seg, การแบ่งส่วนเชิงคลาสสิก) อย่างมีนัยสำคัญ"),
    ("p", "H2 (เชิงการประยุกต์ — การจัดกลุ่ม): ชุด feature 6 กลุ่มที่คำนวณจาก SAM3 mask ผนวกกับ metadata (days_since_last_subculture) โดยเฉพาะสัดส่วนระบบราก (root_ratio) ซึ่งเป็นตัวชี้วัดอันดับ 1 ของความพร้อมอนุบาล สามารถจำแนกขวดออกเป็นกลุ่ม ยังไม่พร้อม / พร้อมอนุบาล / ตรวจเอง ได้ถูกต้อง ≥ 70% เทียบกับการประเมินโดยนักวิทยาศาสตร์ห้องปฏิบัติการ โดยมี minimum sensitivity ≥ 0.6 สำหรับกลุ่มพร้อมอนุบาล"),
],
"ประโยชน์และผลที่คาดว่าจะได้รับ (Benefits and Expected Results)": [
    ("p", "1. เพิ่มประสิทธิภาพห้องปฏิบัติการ: ระบบช่วยคัดกรองขวดที่พร้อมอนุบาลก่อน แล้วให้นักวิทยาศาสตร์ตรวจเฉพาะขวดที่ confidence ต่ำ", {"ind_left": 0.75}),
    ("p", "2. ติดตามแบบไม่ทำลายตัวอย่าง (non-destructive) ไม่ต้องเปิดขวดหรือสัมผัสพืช ลดความเสี่ยงการปนเปื้อน", {"ind_left": 0.75}),
    ("p", "3. เครื่องมือช่วยตัดสินใจ (decision-support) ที่ทำงานข้ามชนิดพืช โดยไม่ต้อง retrain โมเดลต่อชนิด", {"ind_left": 0.75}),
    ("p", "4. จัดลำดับความสำคัญงานในห้องปฏิบัติการ ลดการสูญเสียจาก overcrowding และ hyperhydricity", {"ind_left": 0.75}),
    ("p", "5. ต้นทุนต่ำ ใช้เพียงสมาร์ตโฟน Android + Google Colab + โมเดล open-source และสร้างชุดข้อมูลเปิด (open dataset) ตามหลัก FAIR", {"ind_left": 0.75}),
],
"วัสดุอุปกรณ์และสถานที่ดำเนินงาน (Materials and Workplace/s)": [
    ("p", "รายการวัสดุอุปกรณ์ (List of Materials):", {"bold": True}),
    ("p", "• สมาร์ตโฟน Android (กล้อง 12–50 MP, ใช้ถ่ายภาพ) • ขาตั้งกล้อง + พื้นหลังสีขาว/ดำด้าน (matte) • ตัวอย่างพืชเพาะเลี้ยงเนื้อเยื่อในขวดแก้ว (≥ 100 ขวด พริกจินดา) • บัญชี Google Colab (GPU) • บัญชี Hugging Face + Roboflow ที่มีสิทธิ์เข้าถึงโมเดล SAM3 PCS • ซอฟต์แวร์: Python, OpenCV, ultralytics (YOLO-seg), pandas, matplotlib", {"ind_left": 0.5}),
    ("p", "รายชื่อสถานที่ดำเนินงาน (List of Workplace/s):", {"bold": True}),
    ("p", "• ห้องปฏิบัติการเพาะเลี้ยงเนื้อเยื่อ โรงเรียนวิทยาศาสตร์จุฬาภรณราชวิทยาลัย บุรีรัมย์ (PCSHSBR) • Google Colab (cloud GPU) — ดำเนินการประมวลผลภาพด้วยโมเดล SAM3", {"ind_left": 0.5}),
    ("note", "หมายเหตุ: งานนี้เป็นการศึกษาพืช ไม่มีการวิจัยในมนุษย์หรือสัตว์ทดลอง จึงไม่เข้าข่ายแบบฟอร์ม 4 (Human Participants), 5 (Animal Research) และ 2B (Qualified Scientist)"),
],
"ระเบียบวิธีการทดลอง (Methodology)": [
    ("p", "การเก็บข้อมูล (Data Collection): ถ่ายภาพขวดเพาะเลี้ยงเนื้อเยื่อบนพื้นหลัง matte ด้วยสมาร์ตโฟนระยะคงที่ 20–30 ซม. จัดแสงด้านข้าง 45° หลีกเลี่ยงแสงหน้าเพื่อลด glare ถ่ายซ้ำ 2–3 ครั้งต่อขวด พร้อมบันทึก metadata (วันที่ถ่าย, days_since_last_subculture, ชนิดพืช, การประเมินground truth โดยนักวิทยาศาสตร์)", {"bold": True, "size": 14}),
    ("fig", "รูปที่ 1 สถาปัตยกรรมระบบ (cloud-primary: Android app → Roboflow SAM3 PCS API → on-device feature extraction + decision)"),
    ("p", "ขั้นตอนการประมวลผลภาพ: (1) รวบรวมภาพ → (2) ตรวจจับขอบเขตขวด (bottle ROI) → (3) รัน SAM3 PCS (facebook/sam3) บน GPU แบบ headless batch ด้วยพรอมป์ ['plant','leaf','shoot','stem','root'], score/mask threshold ≥ 0.5 → (4) รับ binary mask ต่อพรอมป์พร้อม confidence และ bbox → (5) นับใบแบบ merged กัน over-segmentation และ fallback นับจาก plant+shoot", {"bold": True, "size": 14}),
    ("p", "การคำนวณ feature 6 กลุ่ม: โครงสร้าง (coverage_ratio, height_proxy, hull_ratio, total_area_px), อวัยวะ (shoot_count, stem_count, root_count, leaf_count), ความซับซ้อน (leaf_merged, leaf_density), สี (green_pct, yellow_ratio, brown_ratio, dark_green_ratio, healthy_color), คุณภาพภาพ (glare_score, condensation_score), และ verdict (readiness_index, confidence)", {"bold": True, "size": 14}),
    ("p", "ขั้นตอนวิธีตัดสินใจ (Decision Algorithm): ใช้ rule-based algorithm ที่อธิบายได้ (interpretable) จัดกลุ่ม 3 คลาส (ยังไม่พร้อม / พร้อมอนุบาล / ตรวจเอง) ด้วยเกณฑ์ coverage_ratio, days_since_last_subculture และ shoot_growth โดยคำนวณ confidence = base × (1 − glare_penalty) และรองรับ manual override เสมอ — หมายเหตุ: เกณฑ์ความพร้อมอนุบาลต้อง calibrate กับข้อมูลจริงและผู้เชี่ยวชาญ ไม่ใช้ coverage สูงเพียงอย่างเดียว (เสี่ยง hyperhydricity)", {"bold": True, "size": 14}),
    ("note", "หมายเหตุ: โมเดล SAM3 เป็น gated model ต้องยืนยันสิทธิ์ผ่าน Hugging Face และใช้ GPU (Colab T4) จึงรันการประมวลผลบนคลาวด์"),
],
"การวิเคราะห์ข้อมูล (Data Analysis)": [
    ("p", "การประเมิน segmentation (เทียบ baseline): คำนวณ mIoU, Dice, precision, recall, F1 ที่ระดับพิกเซล เทียบ mask ของ SAM3 กับ ground truth และเทียบกับ SAM2, YOLO-seg, การแบ่งส่วนเชิงคลาสสิก บนชุดภาพเดียวกัน พร้อมบันทึกเวลาประมวลผลต่อภาพ"),
    ("p", "การประเมินการจัดกลุ่ม: ใช้ confusion matrix, accuracy (เป้าหมาย ≥ 70%), precision/recall ต่อกลุ่ม (minimum sensitivity ≥ 0.6 สำหรับกลุ่มพร้อมอนุบาล) และ F1"),
    ("p", "การวิเคราะห์ความสัมพันธ์: scatter/box plot ของแต่ละ feature จำแนกตามกลุ่ม + correlation matrix เพื่อตรวจ multicollinearity และหาอำนาจจำแนกของ feature; การปรับเทียบ threshold ด้วย validation set (coverage 0.25–0.85 step 0.05, days 14–70 step 7) เลือกชุดที่ให้ MCC สูงสุด"),
],
"ความเสี่ยงและความปลอดภัย (Risk and Safety)": [
    ("p", "ความเสี่ยงด้านเทคนิค: การตรวจจับขวด (ROI) ไม่ชัดในบางภาพ (≈36% ในผลนำร่อง) และแสงสะท้อน/ไอน้ำรบกวนการแบ่งส่วน — แก้โดยการจัดฉากถ่ายมาตรฐาน, กัน ROI ไม่ชัดไม่ให้ verdict ผิด, และใช้ fallback การนับใบ"),
    ("p", "ความเสี่ยงด้านความถูกต้อง: เกณฑ์ความพร้อมอนุบาลยังต้องยืนยันกับผู้เชี่ยวชาญ (ระบบรากเป็นตัวชี้วัดหลัก) และต้อง calibrate ต่อชนิดพืช — ใช้ validation กับ ground truth และ iterative threshold tuning"),
    ("p", "ความปลอดภัยด้านเคมี/ชีวภาพ: อาหารเพาะเลี้ยงเนื้อเยื่อและสารเคมี (สารควบคุมการเจริญเติบโต/Hormone) จัดการตามหลักปฏิบัติของห้องปฏิบัติการตามมาตรฐานความปลอดภัย (risk assessment) และเก็บข้อมูลภาพแบบไม่สัมผัสตัวอย่าง (non-invasive) ลดความเสี่ยงการปนเปื้อน"),
    ("note", "หมายเหตุ: ในโครงการนี้ เนื่องจากเป็นการศึกษาพืชในสภาพเพาะเลี้ยงเนื้อเยื่อแบบไม่ทำลายตัวอย่าง ไม่มีการดำเนินการวิจัยในมนุษย์หรือสัตว์ทดลอง จึงไม่เข้าข่ายข้อกำหนดด้านจริยธรรมการวิจัยในมนุษย์/สัตว์"),
],
"แผนการดำเนินงาน (Research Plan)": [],
"บรรณานุกรม (Bibliography)": [],  # เติมผ่านตาราง/รายการด้านล่าง
}

# จับคู่ heading
head_map = {}
for p in doc.paragraphs:
    t = p.text.strip()
    for h in SECTION_HEADINGS:
        if t == h:
            head_map[h] = p

# 1) clear placeholder หลังทุก heading
for h, hp in head_map.items():
    clear_section_after(hp)

# 2) เติมเนื้อหา
for h, hp in head_map.items():
    if h in C and C[h]:
        add_blocks(hp, C[h])

# 3) เติม บทสรุป (Synopsis) + คำสำคัญ (Keywords) — หา Paragraph แรกหลังนั้น
def set_text_of_para(p, text, bold=False, size=BODY):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    r = p.add_run(text)
    set_font(r, bold=bold, size=size)

for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("คำสำคัญ (Keywords)"):
        set_text_of_para(p, "คำสำคัญ (Keywords): การเพาะเลี้ยงเนื้อเยื่อ; ความพร้อมอนุบาล (acclimatization); SAM3; zero-shot segmentation; non-destructive phenotyping; computer vision")
    elif t.startswith("บทสรุป (Synopsis)") or "Synopsis" in t:
        pass

# เติมบทสรุป: หา paragraph ที่เป็นคำบรรยาย (สีเทา) หลังจาก "บทสรุป" แล้วแทนข้อความ
placed_syn = False
paras = list(doc.paragraphs)
for i, p in enumerate(paras):
    if "Synopsis" in p.text or p.text.strip().startswith("บทสรุป"):
        # หา placeholder ถัดมา (สีเทา) 1-2 อัน แล้วแทน
        for j in range(i+1, min(i+4, len(paras))):
            if paras[j].text.strip() and paras[j].text.strip() == p.text.strip():
                continue
            if is_gray(paras[j]) or paras[j].text.strip().startswith("บรรยาย"):
                set_text_of_para(paras[j], "การเพาะเลี้ยงเนื้อเยื่อพืชเป็นเทคโนโลยีหลักของการขยายพันธุ์พืชเชิงพาณิชย์ แต่การตัดสินใจว่าเมื่อใดจึงควรย้ายต้นออกอนุบาล (acclimatization) ยังคงอาศัยการตรวจด้วยสายตาของนักวิทยาศาสตร์เป็นรายขวด ซึ่งเป็นคอขวดด้านแรงงานและเวลา และมีความแปรปรวนระหว่างผู้ประเมิน งานนี้เสนอระบบคัดกรองความพร้อมอนุบาลของต้นกล้าเพาะเลี้ยงเนื้อเยื่อผ่านขวดแก้วแบบไม่ทำลายตัวอย่าง (non-destructive) ด้วยการแบ่งส่วนภาพ (segmentation) แบบ zero-shot จากโมเดล Segment Anything Model 3 (SAM3) — ปัญหาคอมพิวเตอร์วิทัศน์ที่ท้าทายจากแสงสะท้อน (glare) ไอน้ำ และความโค้งของแก้ว — ร่วมกับการออกแบบพรอมป์ข้อความ 5 คำ (plant, leaf, shoot, stem, root) การตรวจจับขอบเขตขวด (bottle ROI) และขั้นตอนวิธีตัดสินใจแบบกฎที่จัดกลุ่มขวดเป็น ยังไม่พร้อม / พร้อมอนุบาล / ตรวจเอง ระบบนี้ต่อยอดสู่การติดตามการเจริญตามเวลา (time-series monitoring) และสร้างชุดข้อมูลภาพเพาะเลี้ยงเนื้อเยื่อผ่านขวดแก้วชุดแรก โดยประเมินเทียบกับวิธีพื้นฐาน (SAM2, YOLO-seg, การแบ่งส่วนเชิงคลาสสิก) ด้วย mIoU, Dice และ F1 — ต้นทุนต่ำ ใช้เพียงสมาร์ตโฟน")
                placed_syn = True
                break
        break

# 4) เติมชื่อโครงงาน (ปก) — หา "ข้อเสนอโครงงาน (Research Proposal)" และใส่ชื่อใต้
# (ปกเป็นตาราง 0; หายาก ปล่อยให้ผู้ใช้กรอกรหัสโครงการ/ชื่อที่ SIMS)

# 5) เก็บกวาดรอบเดียว: ลบข้อความสีเทา + ส่วนมนุษย์/สัตว์/PHBA/สารเคมี + หัวข้อเปล่า + ขยะ
import re as _re
SPECIAL = [
    "อธิบายช่วงอายุ เพศ", "จะค้นหาหรือรับสมัครผู้เข้าร่วมการวิจัย", "ผู้เข้าร่วมการวิจัยจะถูกให้ทำอะไร",
    "ผู้เข้าร่วมแต่ละคนต้องเข้าร่วมกิจกรรม", "จะลดความเสี่ยงเหล่านี้อย่างไร", "ชื่อ หมายเลขโทรศัพท์ วันเดือนปีเกิด",
    "ข้อมูลจะถูกเก็บเป็นความลับ", "อธิบายผลกระทบหรือประโยชน์ที่คาดว่าจะได้รับจากการวิจัย", "อธิบายสถานที่เลี้ยงสัตว์",
    "อธิบายวิธีการจัดการกับสัตว์เมื่อสิ้นสุดการศึกษา", "อธิบายมาตรการด้านความปลอดภัยอย่างละเอียด",
    "รายละเอียดเพิ่มเติมเกี่ยวกับการทดลองที่มีมนุษย์", "ระบุว่ามีกลุ่มเปราะบาง", "จะใช้แบบสำรวจ แบบสอบถาม",
    "ผู้เข้าร่วมการวิจัยอาจเผชิญความเสี่ยง", "หากเป็นข้อมูลแบบไม่เปิดเผยตัวตน", "ข้อมูลจะถูกจัดเก็บไว้ที่ใด",
    "อธิบายวิธีการแจ้งให้ผู้เข้าร่วมทราบ", "อภิปรายทางเลือกอื่นที่อาจใช้แทนสัตว์ทดลอง", "อธิบายขั้นตอนทั้งหมดที่จะดำเนินการ",
    "ระบุจำนวนสัตว์ ชนิด สายพันธุ์", "รายละเอียดเพิ่มเติมเกี่ยวกับการทดลองที่มีการใช้สัตว์ทดลอง",
    "9. รายละเอียดเพิ่มเติมเกี่ยวกับการทดลองที่มีการใช้สัตว์ทดลอง", "รายละเอียดเพิ่มเติมเกี่ยวกับการทดลองที่มีสารชีวภาพ", "ระบุแหล่งที่มาของสิ่งมีชีวิต",
    "รายละเอียดเพิ่มเติมเกี่ยวกับการทดลองที่มีสารเคมี", "อธิบายกระบวนการประเมินความเสี่ยง", "ต้องศึกษาเอกสารข้อมูลความปลอดภัย",
    "10. รายละเอียดเพิ่มเติมเกี่ยวกับการทดลองที่มีสารชีวภาพ", "11. รายละเอียดเพิ่มเติมเกี่ยวกับการทดลองที่มีสารเคมี",
    "https://commonchemistry.cas.org/", "https://pubchem.ncbi.nlm.nih.gov/", "https://cdxapps.epa.gov/",
    "https://echa.europa.eu/", "บอกถึงคำสำคัญที่สอดคล้อง", "ระบุรายการวัสดุอุปกรณ์ทั้งหมด", "สำหรับรายการสารเคมี",
]

def should_remove(p):
    t = p.text.strip()
    pPr = p._p.find(qn("w:pPr"))
    has_num = pPr is not None and pPr.find(qn("w:numPr")) is not None
    # 1) placeholder numbering ว่าง / เป็นเลขล้วน (2.1 / 2.2)
    if has_num and (not t or (all(c in "0123456789. " for c in t) and "." in t)):
        return True
    if not t:
        return False
    if is_gray(p):
        return True
    if t.startswith("(ใส่ชื่อหัวข้อ)") or t.startswith("(กรุณาลบ") or t.startswith("(ตรงนี้มี"):
        return True
    if "หมายเหตุ: งานนี้เป็นการศึกษาพืช" in t:
        return False  # เก็บ
    if t in ("รายการวัสดุอุปกรณ์ (List of Materials)", "รายชื่อสถานที่ดำเนินงาน (List of Workplace/s)"):
        return True
    for s in SPECIAL:
        if t.startswith(s):
            return True
    if t and all(c in "0123456789. " for c in t) and "." in t:
        return True
    return False

removed = 0
for p in list(doc.paragraphs):
    if should_remove(p):
        el = p._p
        el.getparent().remove(el)
        removed += 1
print("เก็บกวาด (ลบ gray/มนุษย์-สัตว์/สาร/ขยะ):", removed, "| placed synopsis:", placed_syn)

# ลบ note ที่ผมใส่เกิน (# ทั้งนี้ หากเป็น... + (กรุณาลบ...)) — เก็บเฉพาะ continuation note จริง
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("(กรุณาลบข้อความสีเทานี้ทิ้งก่อนนำไปใช้"):
        p._p.getparent().remove(p._p)
        removed += 1

# 6) เติมตาราง Gantt (table 3, 7x10)
for tb in doc.tables:
    txt = tb.rows[0].cells[0].text.strip() if tb.rows else ""
    if "กิจกรรม" in txt:
        # กำหนด month columns: ส.ค. ก.ย. ต.ค. พ.ย. ธ.ค. ม.ค. ก.พ. มี.ค.
        # แผน 3 เดือน: ก.ย.-พ.ย. 2569
        plan = {
            "ค้นหาปัญหาที่สนใจ": {"ก.ย.": "X"},
            "จัดทำข้อเสนอโครงการ": {"ส.ค.": "X", "ก.ย.": "X"},
            "เตรียมข้อมูล/ภาพ": {"ก.ย.": "X"},
            "รันโมเดล/วิเคราะห์": {"ต.ค.": "X", "พ.ย.": "X"},
            "ประเมินผล/เทียบ baseline": {"พ.ย.": "X"},
            "จัดทำรายงาน": {"ธ.ค.": "X"},
        }
        for r_i, row in enumerate(tb.rows[1:], start=1):
            act = row.cells[0].text.strip()
            month_cols = {}
            # หา header month จากแถวแรก
            for c_i, cell in enumerate(tb.rows[0].cells):
                h = cell.text.strip()
                month_cols[c_i] = h
            for c_i, cell in enumerate(row.cells):
                h = month_cols.get(c_i, "")
                marks = plan.get(act, {})
                if any(h.startswith(m) for m in marks):
                    cell.text = "X"
                    for pr in cell.paragraphs:
                        for r in pr.runs:
                            set_font(r, size=12)
        print("Gantt table filled")
    elif "คำศัพท์ที่เกี่ยวข้อง" in (tb.rows[0].cells[0].text if tb.rows else ""):
        # เติม Definitions: คอลัมน์เดียว ใส่รายการคำศัพท์
        defs = ["acclimatization (การอนุบาล/การปรับสภาพ)|การย้ายต้นกล้าออกจากสภาพปลอดเชื้อสู่สภาพแวดล้อมจริง",
                "zero-shot segmentation|การแบ่งส่วนภาพโดยไม่ต้องฝึกโมเดลกับข้อมูลป้ายกำกับของงานนั้น",
                "SAM3 (Segment Anything Model 3)|โมเดล foundation สำหรับการแบ่งส่วนภาพตามพรอมป์",
                "coverage_ratio|สัดส่วนพื้นที่ต้น/ใบ ต่อพื้นที่ขวด (ROI)",
                "root_ratio|สัดส่วนระบบราก — ตัวชี้วัดอันดับ 1 ของความพร้อมอนุบาล"]
        # ตั้ง text ลงแถวหลังจาก column 1 (ที่วางเลข) — ใส่เป็นบรรทัดใน cell แรกแทน XXXXX
        first_cell = tb.rows[0].cells[0]
        for pr in list(first_cell.paragraphs):
            if pr.text.strip() in ("1. XXXXX", "2. XXXXX", "3. XXXXX", "4. XXXXX", "5. XXXXX"):
                continue
        # ง่าย: ลบ XXXXX rows ที่เหลือ + ใส่ข้อความรวมในบล็อกแรก
        tb.rows[0].cells[0].text = ""
        first_cell = tb.rows[0].cells[0]
        first_cell.text = "\n".join("{} — {}".format(i+1, d.replace("|", ": ")) for i, d in enumerate(defs[:5]))
        for pr in first_cell.paragraphs:
            for r in pr.runs:
                set_font(r, size=12)
        # ลบแถวถัดไปที่เป็น XXXXX (ถ้ามี)
        for row in list(tb.rows[1:]):
            if "XXXXX" in row.cells[0].text or "อาจใส่หรือไม่ก็ได้" in row.cells[0].text:
                row._tr.getparent().remove(row._tr)
        print("Definitions table filled")

# 7) ตั้งฟอนต์ body ให้ทั้งเอกสารเป็น Sarabun PSK 16pt (กันคลาด)
def ensure_doc_font():
    # Normal style
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(BODY)
    rPr = st.element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), FONT)

# 8) เพิ่ม section เครื่องมือที่ใช้ & Gen-AI (ก่อนแผนการดำเนินงาน)
import os as _os
LOGO_DIR = os.path.join(HERE, "assets", "logos")

def logo_path(k):
    if not k:
        return None
    p = os.path.join(LOGO_DIR, k + ".png")
    return p if os.path.exists(p) else None

plan_h = None
for p in doc.paragraphs:
    if p.text.strip() == "แผนการดำเนินงาน (Research Plan)":
        plan_h = p; break

if plan_h is not None:
    plan_el = plan_h._p
    # หัวข้อ section
    p_h = make_para("เครื่องมือที่ใช้และการเปิดเผยการใช้ Generative AI (Tools Used & Generative AI Disclosure)", bold=True, size=16)
    plan_el.addprevious(p_h._p)
    anchor = p_h._p
    intro = [
        "ผู้พัฒนาโครงงานใช้เครื่องมือ Generative AI ในหลายขั้นตอนของกระบวนการพัฒนาและเขียนข้อเสนอโครงงาน โดยมีการเปิดเผยรายละเอียดดังนี้ เพื่อความโปร่งใสและสอดคล้องกับแนวปฏิบัติของเวทีการประกวด",
        "ขอบเขตการใช้งาน: (1) การสังเคราะห์วรรณกรรมและตรวจสอบการอ้างอิง — ใช้ Claude, Consensus AI และ PubMed MCP ในการค้นหา/ยืนยันบทความ; (2) การเขียนข้อเสนอ — ใช้ Claude ในการเรียบเรียง/จัดโครงสร้าง; (3) การสร้างโค้ด — ใช้ Claude ในการช่วยเขียนโค้ดต้นแบบ; (4) การสร้างไดอะแกรม/ตาราง",
        "หมายเหตุ: เนื้อหาทั้งหมดได้รับการตรวจทานและปรับแก้โดยผู้พัฒนาโครงงาน และผู้พัฒนาเป็นผู้รับผิดชอบต่อเนื้อหาในข้อเสนอทั้งหมด แม้ส่วนที่เขียนโดยใช้เครื่องมือ Gen-AI ก็ตาม",
    ]
    for t in intro:
        p = make_para(t, size=14)
        anchor = move_after(anchor, p._p)

    # ตารางเครื่องมือ
    tools = [
        ("Claude (Anthropic)", "Claude Opus 4.8", "AI assistant — เขียน/เรียบเรียงข้อเสนอ, ช่วยเขียนโค้ด", "claude"),
        ("Roboflow", "SAM3 PCS API", "zero-shot segmentation (text prompts)", None),
        ("Google Colab", "T4 GPU", "รัน pipeline/batch processing", "colab"),
        ("Hugging Face", "SAM3 model (gated)", "เข้าถึงโมเดล facebook/sam3", "huggingface"),
    ]
    tbl = doc.add_table(rows=len(tools) + 1, cols=3)
    tbl.style = "Table Grid"
    hdr = ["เครื่องมือ", "เวอร์ชัน / บทบาท", "การใช้งาน"]
    # ตั้งค่า cell
    def set_cell(cell, text, bold=False, size=12):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, bold=bold, size=size)
    for c_i, htxt in enumerate(hdr):
        set_cell(tbl.rows[0].cells[c_i], htxt, bold=True)
    for r_i, (name, ver, usage, logo) in enumerate(tools, start=1):
        cells = tbl.rows[r_i].cells
        # col0: logo + ชื่อ
        cells[0].text = ""
        p0 = cells[0].paragraphs[0]
        lp = logo_path(logo)
        if lp:
            p0.add_run().add_picture(lp, width=Cm(0.75))
        r = p0.add_run(" " + name); set_font(r, size=12)
        set_cell(cells[1], ver, size=12)
        set_cell(cells[2], usage, size=12)
    # ย้ายตารางไปวางก่อนแผนงาน (หลังเนื้อหา)
    anchor = move_after(anchor, tbl._tbl)

# 9) เติมบรรณานุกรม (จาก docs/proposal_th_draft.md §13)
import re as _re2
SRC = os.path.join(HERE, "proposal_th_draft.md")
refs = []
if os.path.exists(SRC):
    src_text = open(SRC, encoding="utf-8").read()
    i = src_text.find("## 13. บรรณานุกรม")
    block = ""
    if i != -1:
        rest = src_text[i:]
        j = rest.find("\n---")
        block = rest[:j] if j != -1 else rest
    if block:
        for line in block.splitlines():
            line = line.strip()
            if not line or line.startswith("#") or line.startswith("---") or line == "บรรณานุกรม":
                continue
            # เอาเฉพาะรายการอ้างอิง (มีปี / จุด / DOI)
            if _re2.search(r"\(\d{4}\)|https?://|doi", line):
                refs.append(line)
print("บรรณานุกรม พบ:", len(refs))

if refs:
    bib_par = None
    for p in doc.paragraphs:
        if p.text.strip() == "บรรณานุกรม (Bibliography)" or p.text.strip() == "12. บรรณานุกรม (Bibliography)":
            bib_par = p; break
    if bib_par is not None:
        anchor = bib_par._p
        for ref in refs:
            p = make_para(ref, size=13, ind_left=0.75, after=60)
            # hanging indent
            pf = p.paragraph_format
            from docx.shared import Cm as _Cm
            pf.left_indent = _Cm(1.0)
            pf.first_line_indent = _Cm(-0.5)
            anchor = move_after(anchor, p._p)

ensure_doc_font()
doc.save(OUT)
print("saved:", OUT)
