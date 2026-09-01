# -*- coding: utf-8 -*-
"""build_clone — clone ไฟล์รุ่นพี่ (formatting เป๊ะ) เปลี่ยนเนื้อหาเป็น VitroVision
- ฐาน: editedYSC-Proposal_fractalwall.docx (ล้าง body, เก็บ header/footer/section)
- ฟอนต์ TH SarabunPSK 16pt ทั้งหมด · หัวข้อ bold · body Thai-justify + first-line 0.508 cm
- ตารางปก/คำศัพท์ (1x1 มีขอบ) · ตาราง metrics + แผน · flowchart 3 รูปกลางหน้า
- ไม่ใช้ em dash (—) · footer = เลขหน้ากลาง
"""
import copy
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

SRC = r"C:\Users\User\Downloads\editedYSC-Proposal_fractalwall.docx"
OUT = "docs/VitroVision_Proposal_YSC.docx"
FONT = "TH SarabunPSK"

doc = Document(SRC)
body = doc.element.body

# ---------- ล้าง body (เก็บเฉพาะ sectPr ท้าย) ----------
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)


# ---------- helpers ----------
def set_font(run, bold=False, size=16):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts")
        rPr.insert(0, rF)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(a), FONT)


def para(text, align="j", bold=False, first=0.508, indL=0.0, after=0):
    p = doc.add_paragraph()
    if align == "c":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "r":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    if first:
        pf.first_line_indent = Cm(first)
    if indL:
        pf.left_indent = Cm(indL)
    if after:
        pf.space_after = Pt(after if after else 4)
    pf.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold)
    return p


def bullet(text):
    para("• " + text, first=0.508)


def image(path, width=15.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Cm(13.2))
    return p


def _cell_borders(tc):
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "000000")
        b.append(e)
    tcPr.append(b)


def table(rows, header=True, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            _cell_borders(cell._tc)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (header and ri == 0) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_font(r, bold=(header and ri == 0), size=16)
    if widths:
        for ci, w in enumerate(widths):
            for ri in range(len(rows)):
                t.cell(ri, ci).width = Cm(w)
    return t


def box_table(lines):
    """ตาราง 1x1 มีขอบ (ปก / คำศัพท์) — lines = list of (text, bold, center)"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _cell_borders(cell._tc)
    first = True
    for text, bold, center in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        if text:
            r = p.add_run(text)
            set_font(r, bold=bold)
    return t


# ============================================================ เนื้อหา
para("รหัสโครงการ 29YCSE00054T", align="r", bold=True, first=0)
para("", first=0)

box_table([
    ("ข้อเสนอโครงงาน (Research Proposal)", True, True),
    ("การประกวดโครงงานของนักวิทยาศาสตร์รุ่นเยาว์ ครั้งที่ 29 (YSC 2027)", True, True),
    ("", False, False),
    ("ชื่อโครงงาน (ไทย) VitroVision : การประยุกต์ใช้ปัญญาประดิษฐ์เชิงคอมพิวเตอร์วิทัศน์เพื่อวิเคราะห์และทำนายการเจริญเติบโตของพืชเพาะเลี้ยงเนื้อเยื่อ", True, False),
    ("(อังกฤษ) VitroVision : Application of AI-Based Computer Vision for Analyzing and Predicting the Growth of Tissue Cultured Plants", True, False),
    ("โครงงานสาขา วิทยาศาสตร์คอมพิวเตอร์และปัญญาประดิษฐ์ (Computer Science and Artificial Intelligence, CSAI)", False, False),
    ("ผู้พัฒนา นายพีรดนย์ ด้วงทอง ชั้นมัธยมศึกษาปีที่ ____", False, False),
    ("สถาบันการศึกษา โรงเรียนวิทยาศาสตร์จุฬาภรณราชวิทยาลัย บุรีรัมย์", False, False),
    ("อาจารย์ที่ปรึกษา ________________________________________", False, False),
])
para("", first=0)

para("บทคัดย่อ/บทสรุป (Abstract/Synopsis)", align="c", bold=True, first=0)
para("", first=0)
para("การเพาะเลี้ยงเนื้อเยื่อพืชเป็นเทคโนโลยีการขยายพันธุ์พืชเชิงพาณิชย์ที่สำคัญ แต่การประเมินการเจริญและการตัดสินใจว่าต้นกล้าพร้อมย้ายออกอนุบาลเมื่อใด ยังอาศัยการตรวจด้วยสายตารายขวด ซึ่งใช้แรงงาน เวลา และมีความแปรปรวนระหว่างผู้ประเมิน การเปิดขวดเพื่อวัดมีความเสี่ยงต่อการปนเปื้อน จึงจำเป็นต้องประเมินแบบไม่ทำลายผ่านขวดแก้วปิด ซึ่งมีอุปสรรคจากแสงสะท้อน ไอน้ำ และความโค้งของแก้ว")
para("โครงงานนี้จึงมุ่งพัฒนาและประเมินระบบวิเคราะห์การเจริญเติบโตของพืชเพาะเลี้ยงเนื้อเยื่อแบบไม่ทำลายผ่านขวดแก้วปิดด้วยคอมพิวเตอร์วิทัศน์ โดยเป็นงานต่อเนื่องที่เริ่มดำเนินการแล้วและจะแล้วเสร็จภายในปีนี้ งานที่เริ่มไปแล้วประกอบด้วย (1) การทบทวนวรรณกรรมด้านการแบ่งส่วนภาพสำหรับงานเพาะเลี้ยงเนื้อเยื่อและแบบจำลองพื้นฐาน (2) การสร้างชุดข้อมูลภาพถ่ายจริงผ่านขวดแก้ว 100 ขวด พร้อมค่าอ้างอิงจากผู้ประเมิน และ (3) การทดลองนำร่องด้วยแบบจำลอง Segment Anything Model 3 (SAM3) แบบ zero-shot ผลเบื้องต้นชี้ความเป็นไปได้ของแนวทาง โดยค่าลักษณะที่วัดจากภาพสอดคล้องกับการวัดด้วยมือ (สหสัมพันธ์ของความสูง r = 0.638 ดีกว่า classical 0.498 และ YOLO-COCO 0.133) และเกณฑ์ความสูงของต้นสอดคล้องกับการประเมินของผู้เชี่ยวชาญในเบื้องต้น (accuracy 0.755, sensitivity 0.917) ทั้งนี้ผลดังกล่าวยังเป็นเพียงผลนำร่อง ต้องยืนยันด้วยการประเมินที่เข้มข้นขึ้น")
para("แผนดำเนินงานต่อเนื่องที่เหลือครอบคลุม การประเมินเปรียบเทียบกับวิธีพื้นฐาน (SAM2, YOLO-seg, การแบ่งส่วนเชิงคลาสสิก) การตรวจสอบระดับพิกเซลกับ ground truth ที่ผู้ประเมินกำกับ (mIoU, Dice, F1) การวิเคราะห์ความไวของพรอมป์และเกณฑ์ตัดสินใจ การกลั่นแบบจำลองขนาดใหญ่เป็น U-Net ขนาดเล็กที่ทำงานบนอุปกรณ์ทั่วไป และการทดสอบข้ามชนิดพืช เพื่อให้ได้ระบบต้นทุนต่ำที่ห้องปฏิบัติการเพาะเลี้ยงเนื้อเยื่อใช้งานได้จริง พร้อมชุดข้อมูลเปิดสำหรับการวิจัยต่อยอด")
para("คำสำคัญ (Keywords): การเพาะเลี้ยงเนื้อเยื่อ, การวัดลักษณะปรากฏของพืชด้วยภาพ, การแบ่งส่วนภาพ, แบบจำลองพื้นฐาน, zero-shot, การทำนายการเจริญ")

para("บทนำ (Introduction)", bold=True, first=0)
para("การเพาะเลี้ยงเนื้อเยื่อพืช (plant tissue culture หรือ micropropagation) เป็นเทคโนโลยีการขยายพันธุ์พืชที่ใช้ชิ้นส่วนขนาดเล็ก (explant) เพาะเลี้ยงในสภาพปลอดเชื้อบนอาหารสังเคราะห์ และถูกใช้อย่างกว้างขวางในเชิงพาณิชย์ครอบคลุมพืชเกษตร อาหาร เภสัชกรรม และเครื่องสำอางทั่วโลก (Hasnain et al., 2022; Chandran et al., 2020) สำหรับประเทศไทย การเพาะเลี้ยงเนื้อเยื่อมีบทบาทสำคัญในอุตสาหกรรมกล้วยไม้ซึ่งเป็นสินค้าส่งออกสำคัญ (Thammasiri, 2015) และหน่วยงานวิจัย เช่น ไบโอเทค สวทช. ได้พัฒนาเทคโนโลยีนี้เชิงพาณิชย์สำหรับพืชหลายชนิด")
para("อย่างไรก็ตาม กระบวนการหนึ่งที่ยังคงพึ่งพาแรงงานคนอย่างมากคือการตัดสินใจว่าต้นกล้าพร้อมย้ายออกอนุบาล (acclimatization) เมื่อใด ซึ่งต้องพิจารณารายขวดทุก 3 ถึง 8 สัปดาห์ขึ้นกับชนิดพืช (Pastelín Solano et al., 2019; Regni et al., 2025) การตัดสินใจผิดพลาด ไม่ว่าจะย้ายเร็วหรือช้าเกินไป อาจทำให้เนื้อเยื่อตาย (necrosis) และลดประสิทธิภาพการขยายพันธุ์ (Abdalla et al., 2022) การตรวจด้วยสายตามีความแปรปรวนระหว่างผู้ประเมิน และในห้องปฏิบัติการที่มีขวดหลายร้อยถึงหลายพันขวด การตรวจละเอียดรายขวดทำได้ไม่ทั่วถึง")
para("งานวิจัยก่อนหน้านำคอมพิวเตอร์วิทัศน์มาใช้วัดพืชในขวดแบบไม่ทำลาย เช่น ระบบ Phenomenon ที่ใช้ multi-sensor และ random forest วัด projected area และ canopy height (Bethge et al., 2023) และงานของ Regni et al. (2025) ที่ใช้ภาพ 3D จากสมาร์ตโฟนวัด canopy และ shoot density ใน blackberry และ blueberry แต่ระบบเหล่านี้ใช้ฮาร์ดแวร์เฉพาะราคาสูง หรือยังไม่ใช้แบบจำลองพื้นฐาน (foundation model) ที่ทำงานแบบ zero-shot ข้ามชนิดพืชได้ ช่องว่างที่โครงงานนี้มุ่งตอบคือ (1) ยังไม่มีระบบต้นทุนต่ำที่ติดตามการเจริญของต้นกล้าในขวดตามเวลาแบบครบวงจร (2) ยังไม่มีงานที่ใช้ zero-shot foundation model แบ่งส่วนต้นผ่านขวดแก้วโดยตรงและประเมินเทียบกับวิธีพื้นฐาน และ (3) ยังไม่มีชุดข้อมูลเปิดของภาพเพาะเลี้ยงเนื้อเยื่อผ่านขวดแก้วสำหรับการวิจัยต่อยอด")

box_table([
    ("คำศัพท์ที่เกี่ยวข้อง (Definitions)", True, True),
    ("SAM3 (Segment Anything Model 3) คือ แบบจำลองพื้นฐานสำหรับการแบ่งส่วนภาพที่รับพรอมป์ข้อความและทำงานได้โดยไม่ต้องฝึกโมเดลใหม่ (zero-shot)", False, False),
    ("Ground truth คือ ข้อมูลอ้างอิงที่กำกับโดยมนุษย์ เช่น การประเมินความพร้อมของผู้เชี่ยวชาญ หรือ mask ที่วาดด้วยมือ ใช้เป็นมาตรฐานในการตรวจสอบความถูกต้องของระบบ", False, False),
    ("mIoU (Mean Intersection over Union) คือ ตัวชี้วัดมาตรฐานของความแม่นยำการแบ่งส่วนภาพ วัดความซ้อนทับระหว่าง mask ที่ระบบทำนายกับ ground truth", False, False),
    ("Zero-shot segmentation คือ การแบ่งส่วนภาพโดยไม่ต้องฝึกโมเดลกับข้อมูลชนิดนั้นเพิ่มเติม อาศัยการชี้นำด้วยพรอมป์ข้อความ", False, False),
])

para("2. วัตถุประสงค์ (Objective/s)", bold=True, first=0)
para("\t2.1 เพื่อพัฒนาโมเดลคอมพิวเตอร์วิทัศน์สำหรับวิเคราะห์และวัดลักษณะการเจริญเติบโตของต้นเพาะเลี้ยงเนื้อเยื่อในขวดแก้วแบบไม่ทำลายตัวอย่าง")
para("\t2.2 เพื่อทดสอบและเปรียบเทียบประสิทธิภาพของระบบกับวิธีพื้นฐาน (SAM2, YOLO-seg, การแบ่งส่วนเชิงคลาสสิก) และกับค่าอ้างอิงจากผู้ประเมิน ด้วยตัวชี้วัดมาตรฐาน (mIoU, Dice, F1, accuracy, sensitivity)")
para("\t2.3 เพื่อแสดงให้เห็นว่าระบบช่วยลดภาระงานและเวลาในการเก็บข้อมูลและวิเคราะห์ข้อมูลในห้องปฏิบัติการเพาะเลี้ยงเนื้อเยื่อได้จริง")

para("3. สมมติฐาน (Hypothesis/es)", bold=True, first=0)
para("H1 (ด้านการแบ่งส่วนภาพ): SAM3 ที่ใช้พรอมป์ข้อความ 5 คำ (plant, leaf, shoot, stem, root) สามารถแบ่งส่วนต้นพืชเพาะเลี้ยงเนื้อเยื่อผ่านขวดแก้วที่มี glare ไอน้ำ และความโค้งของแก้วได้ โดยมีค่า mIoU เฉลี่ยไม่ต่ำกว่า 0.65 เมื่อเทียบกับ ground truth ที่ผู้ประเมินกำกับ และสูงกว่าวิธีพื้นฐานอย่างมีนัยสำคัญ")
para("H2 (ด้านการประยุกต์): ชุด feature เชิงปริมาณที่คำนวณจาก mask โดยเฉพาะความสูงของต้น (height_proxy) สามารถจัดกลุ่มความพร้อมอนุบาล (ยังไม่พร้อม / พร้อมอนุบาล) ได้ถูกต้องไม่ต่ำกว่าร้อยละ 70 เทียบกับการประเมินของนักวิทยาศาสตร์ในห้องปฏิบัติการ โดยมีความไว (sensitivity) ไม่ต่ำกว่า 0.6 สำหรับกลุ่มพร้อมอนุบาล ผลนำร่องเบื้องต้นชี้ทิศทางสอดคล้อง (accuracy 0.755, sensitivity 0.917) และยังต้องยืนยันเพิ่มเติม")

para("4. วัสดุอุปกรณ์และสถานที่ดำเนินงาน (Materials and Workplace/s)", bold=True, first=0)
para("4.1 รายการวัสดุอุปกรณ์ (List of Materials)", bold=True, first=0.508)
para("4.1.1 สมาร์ตโฟน (กล้อง 12 ถึง 50 MP) และขาตั้งกล้องพร้อมพื้นหลังด้าน (matte) สำหรับถ่ายภาพขวดในระยะและแสงคงที่")
para("4.1.2 ขวดเพาะเลี้ยงเนื้อเยื่อมาตรฐาน (ขวดแก้วใส) และพืชเพาะเลี้ยงเนื้อเยื่ออย่างน้อย 2 ถึง 3 ชนิดจากห้องปฏิบัติการ")
para("4.1.3 บัญชี Google Colab (GPU ฟรี) และสิทธิ์เข้าถึงโมเดล facebook/sam3 (Hugging Face, gated)")
para("4.1.4 คอมพิวเตอร์สำหรับพัฒนา (เชื่อมต่ออินเทอร์เน็ต) พร้อมซอฟต์แวร์ Python, OpenCV, PyTorch, segmentation-models-pytorch")
para("4.2 รายชื่อสถานที่ดำเนินงาน (List of Workplace/s)", bold=True, first=0.508)
para("4.2.1 ห้องปฏิบัติการเพาะเลี้ยงเนื้อเยื่อพืช โรงเรียนวิทยาศาสตร์จุฬาภรณราชวิทยาลัย บุรีรัมย์ ใช้ถ่ายภาพชุดข้อมูลและเก็บค่าอ้างอิงจากผู้ประเมิน")
para("4.2.2 Google Colab และ Hugging Face ใช้ประมวลผลโมเดล เทรนโมเดล และเผยแพร่ผลงานออนไลน์")

para("5. ระเบียบวิธีการทดลอง (Methodology)", bold=True, first=0)
para("5.1 ภาพรวมการทำงานของระบบ", bold=True, first=0.508)
image("docs/assets/flow_overview.png")
para("ระบบทำงานเป็น 5 ขั้นตอนหลัก ได้แก่ ถ่ายภาพขวด ตรวจจับขอบเขตขวด แบ่งส่วนภาพด้วย SAM3 คำนวณ feature และตัดสินใจด้วยกฎ ภาพที่ประมวลผลไม่ชัด (มี glare ฝ้า หรือไม่พบขวด) จะถูกส่งให้ผู้เชี่ยวชาญตรวจแทนการตัดสินใจอัตโนมัติเพื่อป้องกันความผิดพลาด")
para("5.2 การเก็บข้อมูล (Data Collection)", bold=True, first=0.508)
bullet("ถ่ายภาพขวดบนพื้นหลังด้าน (matte) ด้วยสมาร์ตโฟนระยะคงที่ 20 ถึง 30 เซนติเมตร จัดแสงด้านข้างมุมประมาณ 45 องศา หลีกเลี่ยงแสงตรงเพื่อลด glare และถ่ายซ้ำ 2 ถึง 3 ครั้งต่อขวด")
bullet("บันทึก metadata คู่กับภาพ ได้แก่ วันที่ถ่าย จำนวนวันหลังตัดย้ายครั้งล่าสุด (days_since_last_subculture) ชนิดพืช และการประเมินโดยนักวิทยาศาสตร์ (ground truth: ยังไม่พร้อม / พร้อมอนุบาล / ตรวจเอง)")
bullet("เป้าหมายตัวอย่าง ไม่น้อยกว่า 100 ขวด กระจายครอบคลุม 3 คลาส และครอบคลุมพืชอย่างน้อย 2 ถึง 3 ชนิด")
para("5.3 ขั้นตอนการประมวลผลภาพ (Image Processing Pipeline)", bold=True, first=0.508)
image("docs/assets/flow_pipeline.png")
para("1. รวบรวมภาพถ่ายขวดจากชุดข้อมูล")
para("2. ตรวจจับขอบเขตขวด (bottle ROI detection) เพื่อใช้เป็นกรอบอ้างอิงของพื้นที่ปกคลุม")
para("3. รัน SAM3 PCS (facebook/sam3) บน GPU แบบ headless batch ด้วยพรอมป์ข้อความ 5 คำ (plant, leaf, shoot, stem, root) เกณฑ์ score และ mask threshold 0.5 ขึ้นไป")
para("4. รับผลลัพธ์เป็น binary mask ต่อพรอมป์ พร้อม confidence score และ bounding box")
para("5. นับใบแบบ merged (รวมชิ้นส่วนที่ติดกันเป็น 1 ใบเพื่อลด over-segmentation) พร้อม fallback นับจาก plant และ shoot")
para("5.4 การคำนวณ feature และเกณฑ์ตัดสินใจ", bold=True, first=0.508)
para("ระบบคำนวณ feature เชิงปริมาณ 6 กลุ่มจาก mask ได้แก่ โครงสร้าง (coverage_ratio, height_proxy, hull_ratio) อวัยวะ (leaf_count, shoot_count, root_count) ความซับซ้อน (compactness) สี (green_pct, yellow_ratio, brown_ratio) คุณภาพภาพ (glare_score, condensation_score) และการตัดสินใจ (verdict, confidence) เกณฑ์การจัดกลุ่มใช้ rule-based algorithm ที่อธิบายได้ โดยต้นที่มี height_proxy ตั้งแต่ 0.275 จัดเป็น พร้อมอนุบาล ต่ำกว่านั้นจัดเป็น ยังไม่พร้อม (อ้างอิงจากผลนำร่องที่พบว่าความสูงเป็นตัวชี้วัดที่สอดคล้องกับผู้เชี่ยวชาญมากกว่าความแน่นของทรงพุ่ม) ภาพที่ประมวลผลไม่ชัดถูกทำเครื่องหมาย ตรวจเอง เพื่อส่งให้มนุษย์ตรวจ")
para("5.5 การตรวจสอบความถูกต้อง (Validation)", bold=True, first=0.508)
para("1. ระดับพิกเซล สร้าง ground-truth masks (ผู้ประเมินกำกับ) ไม่น้อยกว่า 30 ภาพ แล้วเทียบ mask ของระบบและวิธีพื้นฐานด้วย mIoU, Dice, F1, precision, recall")
para("2. ระดับการจัดกลุ่ม เทียบ verdict ของระบบกับการประเมินของผู้เชี่ยวชาญด้วย confusion matrix และตัวชี้วัด accuracy, precision, sensitivity, specificity, F1, MCC, Cohen's kappa")
para("3. ความสอดคล้องระหว่างผู้ประเมิน ประเมิน inter-rater reliability (ICC และ Cohen's kappa) เพื่อเทียบความแปรปรวนของระบบกับความแปรปรวนของมนุษย์เอง")
para("4. การทดสอบข้ามชนิด ทดสอบว่าระบบทำงานกับพืชต่างชนิดได้โดยไม่ต้องฝึกโมเดลใหม่")

para("6. การวิเคราะห์ข้อมูล (Data Analysis)", bold=True, first=0)
para("การประเมินผลใช้ตัวชี้วัดมาตรฐานที่ใช้กันทั่วไปในงาน benchmark ของแบบจำลอง AI แบ่งเป็น 2 ระดับ")
para("6.1 การประเมินการแบ่งส่วนภาพ (Segmentation ระดับพิกเซล)", bold=True, first=0.508)
table([
    ["ตัวชี้วัด", "นิยาม", "ความหมาย"],
    ["mIoU (Mean Intersection over Union)", "TP / (TP+FP+FN) เฉลี่ยทุกพิกเซล", "มาตรฐานหลักของการแบ่งส่วนภาพ"],
    ["Dice (F1 ระดับพิกเซล)", "2TP / (2TP+FP+FN)", "ความซ้อนทับของ mask"],
    ["Precision / Recall (พิกเซล)", "TP/(TP+FP) และ TP/(TP+FN)", "ความแม่นยำและความครบถ้วนของพื้นที่"],
    ["Pixel Accuracy", "(TP+TN) / ทั้งหมด", "สัดส่วนพิกเซลที่จำแนกถูก"],
], widths=[5.2, 5.2, 6.0])
para("เทียบ mask ของระบบ (SAM3 และ U-Net) กับ ground truth ของผู้ประเมิน และเทียบกับวิธีพื้นฐาน (SAM2, YOLO-seg, การแบ่งส่วนเชิงคลาสสิก) บนชุดภาพเดียวกัน พร้อมบันทึกเวลาประมวลผลต่อภาพเพื่อพิจารณาความคุ้มค่าระหว่างความแม่นยำกับต้นทุนการคำนวณ")
para("6.2 การประเมินการจัดกลุ่มความพร้อม (Classification ระดับภาพ)", bold=True, first=0.508)
table([
    ["ตัวชี้วัด", "นิยาม", "เป้าหมาย"],
    ["Accuracy", "(TP+TN) / ทั้งหมด", "ไม่ต่ำกว่าร้อยละ 70"],
    ["Precision (ต่อกลุ่ม)", "TP / (TP+FP)", "รายงานทุกกลุ่ม"],
    ["Sensitivity/Recall (ต่อกลุ่ม)", "TP / (TP+FN)", "ไม่ต่ำกว่า 0.6 (กลุ่มพร้อมอนุบาล)"],
    ["Specificity", "TN / (TN+FP)", "รายงานทุกกลุ่ม"],
    ["F1-score", "2·P·R / (P+R)", "รายงานทุกกลุ่ม"],
    ["MCC (Matthews Correlation Coefficient)", "ครอบคลุมทั้ง 4 ช่องของ confusion matrix", "ใช้เลือก threshold (ทนต่อคลาสไม่สมดุล)"],
    ["Cohen's kappa", "ความสอดคล้องหักความบังเอิญ", "เทียบระบบกับผู้ประเมิน"],
], widths=[5.2, 6.0, 5.2])
para("เนื่องจากชุดข้อมูลมีคลาสไม่สมดุล (imbalance) การรายงานจึงไม่ใช้ accuracy เพียงตัวเดียว แต่ใช้ precision, sensitivity, F1, MCC และ Cohen's kappa ร่วมกัน เพื่อให้เห็นภาพครบทุกมิติของความผิดพลาด")
para("6.3 การปรับเทียบเกณฑ์ (Threshold Tuning) และการวิเคราะห์ความไว", bold=True, first=0.508)
bullet("ทดสอบเกณฑ์ความพร้อมช่วงกว้าง (เช่น 0.10 ถึง 0.90 ทีละ 0.05) แล้วเลือกค่าที่ให้ MCC สูงสุด ซึ่งเหมาะกับคลาสไม่สมดุลมากกว่า accuracy")
bullet("วิเคราะห์ความไวของพรอมป์ (prompt sensitivity) โดยทดสอบชุดพรอมป์ทางเลือกแล้วรายงานความแปรปรวนของ mIoU และ verdict (อ้างอิงงาน Dubois et al., 2026 ที่พบว่า SAM3 แบบชี้นำด้วยข้อความไวต่อถ้อยคำพรอมป์)")
bullet("วิเคราะห์ความสัมพันธ์ระหว่าง feature กับความพร้อมอนุบาลด้วย scatter plot, box plot และ correlation matrix เพื่อตรวจ multicollinearity และหาอำนาจจำแนกของแต่ละ feature")
para("6.4 ผังการตรวจสอบความถูกต้อง", bold=True, first=0.508)
image("docs/assets/flow_validation.png")

para("7. แผนการดำเนินงาน (Research Plan)", bold=True, first=0)
table([
    ["กิจกรรม", "ก.ย. 69", "ต.ค. 69", "พ.ย. 69", "ธ.ค. 69", "ผู้รับผิดชอบ"],
    ["ส่งข้อเสนอ YSC 2027 (ภายใน 10 ก.ย.)", "●", "", "", "", "พีรดนย์ ด้วงทอง"],
    ["เก็บข้อมูลภาพถ่ายซ้ำรายรอบ + metadata (time-series)", "●", "●", "", "", "พีรดนย์ ด้วงทอง"],
    ["สร้าง ground-truth masks (ไม่น้อยกว่า 30 ภาพ) + วัดมือ", "●", "●", "", "", "พีรดนย์ ด้วงทอง"],
    ["รัน baseline (SAM2 / YOLO-seg / classical) + mIoU/Dice", "", "●", "", "", "พีรดนย์ ด้วงทอง"],
    ["เทรน U-Net ขนาดเล็ก (กลั่นจาก SAM3) + ทดสอบบนชุด 100 ขวด", "", "●", "●", "", "พีรดนย์ ด้วงทอง"],
    ["วิเคราะห์ growth curve + threshold sensitivity (MCC สูงสุด)", "", "", "●", "", "พีรดนย์ ด้วงทอง"],
    ["ตรวจข้ามชนิด (cross-species) + inter-rater", "", "", "●", "", "พีรดนย์ ด้วงทอง"],
    ["เขียนรายงานฉบับสมบูรณ์ + เปิดชุดข้อมูล (FAIR)", "", "", "", "●", "พีรดนย์ ด้วงทอง"],
], widths=[7.0, 1.5, 1.5, 1.5, 1.5, 3.0])
para("หมายเหตุ: แผนยืดหยุ่นได้ตามความพร้อมของวัสดุพืชในห้องปฏิบัติการและผลการทดสอบระหว่างทาง ครอบคลุมการดำเนินงานต่อเนื่องจนจบปี 2569 ตามกรอบเวลาการประกวด", first=0.508)

para("8. ประโยชน์และผลที่คาดว่าจะได้รับ (Benefits and Expected Results)", bold=True, first=0)
para("8.1 ช่วยเพิ่มประสิทธิภาพห้องปฏิบัติการเพาะเลี้ยงเนื้อเยื่อ โดยระบบคัดกรองขวดที่พร้อมอนุบาลก่อน ลดเวลาที่นักวิทยาศาสตร์ต้องตรวจขวดทีละขวด")
para("8.2 ติดตามการเจริญแบบไม่ทำลายตัวอย่าง ไม่ต้องเปิดขวดหรือสัมผัสพืช ลดความเสี่ยงการปนเปื้อน (สอดคล้องกับ Bethge et al., 2023)")
para("8.3 เครื่องมือช่วยตัดสินใจที่ทำงานข้ามชนิดพืชโดยไม่ต้องฝึกโมเดลใหม่ต่อชนิด ช่วยลดต้นทุนการพัฒนาโมเดลเฉพาะพืช")
para("8.4 บันทึกประวัติการเจริญตามเวลา (growth history) ข้อมูลอนุกรมเวลาต่อขวดใช้วิเคราะห์แนวโน้มและคาดการณ์ล่วงหน้าได้")
para("8.5 ต้นทุนต่ำ ใช้เพียงสมาร์ตโฟนและบริการ GPU ฟรี ไม่ต้องซื้อฮาร์ดแวร์เฉพาะ")
para("8.6 ชุดข้อมูลเปิด (open dataset) ตามหลัก FAIR เป็นทรัพยากรใหม่สำหรับงาน phenotyping ในสภาพเพาะเลี้ยงเนื้อเยื่อที่ยังขาดแคลน")
para("8.7 ต้นแบบของแบบจำลอง U-Net ขนาดเล็ก (กลั่นจาก SAM3) ที่ทำงานบนอุปกรณ์ทั่วไป ต่อยอดเป็นแอปพลิเคชันมือถือสำหรับใช้งานจริงในห้องปฏิบัติการ")

para("9. การเปิดเผยข้อมูลเกี่ยวกับ Generative AI และเทคโนโลยีปัญญาประดิษฐ์ที่ช่วยในกระบวนการจัดทำข้อเสนอ (Disclosure of Generative AI and AI-Assisted Technologies in the Writing Process)", bold=True, first=0)
para("โครงงานนี้ใช้เครื่องมือ Generative AI (Gen-AI) ในบางขั้นตอนของกระบวนการพัฒนาและเขียนข้อเสนอ โดยเปิดเผยรายละเอียดดังนี้")
para("9.1 เครื่องมือที่ใช้", bold=True, first=0.508)
table([
    ["เครื่องมือ", "ผู้พัฒนา", "การใช้งาน"],
    ["DeepSeek", "DeepSeek", "ช่วยสังเคราะห์วรรณกรรม เรียบเรียงและตรวจทานข้อความ ช่วยเขียนโค้ดต้นแบบ และออกแบบโครงสร้างเอกสาร"],
], widths=[4.0, 4.0, 8.0])
para("9.2 ขอบเขตการใช้งาน", bold=True, first=0.508)
para("1. การสังเคราะห์วรรณกรรม ใช้ DeepSeek ช่วยสรุปสาระจากบทความวิจัยที่ผู้พัฒนาคัดเลือกมา โดยผู้พัฒนาเป็นผู้ตรวจยืนยันการอ้างอิงทุกจุดจากแหล่งต้นทางจริง (DOI/URL ที่เข้าถึงได้) และไม่มีรายการอ้างอิงใดที่มาจากการสร้างข้อมูลเทียมของ AI (hallucination)")
para("2. การเขียนข้อเสนอ ใช้ DeepSeek ช่วยเรียบเรียงเนื้อหา เช่น การจัดโครงสร้างบทนำ วิธีดำเนินการ และภาษาไทยทางวิชาการ โดยเนื้อหาทั้งหมดได้รับการตรวจทานและปรับแก้โดยผู้พัฒนาก่อนนำไปใช้")
para("3. การสร้างโค้ด ใช้ DeepSeek ช่วยเขียนโค้ดต้นแบบสำหรับการประมวลผลภาพ การเชื่อมต่อโมเดล และการคำนวณ feature โดยโค้ดทุกส่วนผ่านการตรวจสอบและทดสอบโดยผู้พัฒนา")
para("4. การสร้างแผนภาพ ใช้ Mermaid และ matplotlib (Python) ในการสร้าง flowchart และแผนภาพ โดยผู้พัฒนาเป็นผู้กำหนดโครงสร้างและตรวจทานเนื้อหาของแผนภาพเอง และไม่ได้ใช้เครื่องมือสร้างภาพด้วย Gen-AI (image-generation) สำหรับแผนภาพ")
para("9.3 ข้อจำกัดและการตรวจสอบโดยมนุษย์", bold=True, first=0.508)
para("ผู้พัฒนาตระหนักดีว่า Generative AI มีข้อจำกัด ได้แก่ อาจสร้างข้อมูลอ้างอิงที่ไม่มีอยู่จริง (hallucination) ความไม่แม่นยำของเนื้อหาเชิงเทคนิคเฉพาะทาง และอคติที่แฝงในข้อมูลเทรนของโมเดล ดังนั้นทุกเนื้อหาที่ได้จาก Gen-AI จึงผ่านการตรวจสอบโดยผู้พัฒนา โดยการอ้างอิงตรวจเทียบแหล่งต้นทางจริง โค้ดทดสอบการทำงานจริง เนื้อหาชีววิทยาตรวจกับเอกสารปฐมภูมิ และภาษากับข้อความตรวจทานโดยผู้พัฒนา ผู้พัฒนาโครงงานขอรับผิดชอบต่อเนื้อหาทั้งหมดในข้อเสนอ แม้ส่วนที่เขียนโดยใช้เครื่องมือ Gen-AI ก็ตาม")

para("10. บรรณานุกรม (Bibliography)", bold=True, first=0)
bib = [
    "Abdalla, N., El-Ramady, H., Seliem, M. K., El-Mahrouk, M. E., Taha, N., Bayoumi, Y., Shalaby, T. A., & Dobránszki, J. (2022). An academic and technical overview on plant micropropagation challenges. Horticulturae, 8(8), 677. https://doi.org/10.3390/horticulturae8080677",
    "Barua, K. N., Singha, B. L., Bordoloi, S., & Bora, B. (2022). In vitro seed propagation and mass multiplication of some magnificent orchids of Northeast India. Journal of Medicinal Plants Studies, 10(2c), 208–213. https://doi.org/10.22271/plants.2022.v10.i2c.1411",
    "Bethge, H., Winkelmann, T., Lüdeke, P., & Rath, T. (2023). Low-cost and automated phenotyping system \"Phenomenon\" for multi-sensor in situ monitoring in plant in vitro culture. Plant Methods, 19, 42. https://doi.org/10.1186/s13007-023-01018-w",
    "Carion, N., Gustafson, L., Hu, Y.-T., Debnath, S., Hu, R., Suris, D., Ryali, C., et al. (2025). SAM 3: Segment anything with concepts. arXiv. https://arxiv.org/abs/2511.16719",
    "Chandran, H., Meena, M., Barupal, T., & Sharma, K. (2020). Plant tissue culture as a perpetual source for production of industrially important bioactive compounds. Biotechnology Reports, 26, e00450. https://doi.org/10.1016/j.btre.2020.e00450",
    "Dubois, R., Bousset, L., Jumel, S., Leclerc, M., Parisey, N., & Joly, A. (2026). Text guidance is powerful but prompt-sensitive for weakly-supervised leaf symptom segmentation (preprint). bioRxiv. https://doi.org/10.64898/2026.07.10.737680",
    "Hasnain, A., Naqvi, S. A. H., Ayesha, S. I., Khalid, F., Ellahi, M., Iqbal, S., et al. (2022). Plants in vitro propagation with its applications in food, pharmaceuticals and cosmetic industries; current scenario and future approaches. Frontiers in Plant Science, 13, 1009395. https://doi.org/10.3389/fpls.2022.1009395",
    "Kirillov, A., Mintun, E., Ravi, N., Mao, H., Rolland, C., Gustafson, L., Xiao, T., et al. (2023). Segment anything. Proceedings of the IEEE/CVF International Conference on Computer Vision (ICCV) 2023. https://arxiv.org/abs/2304.02643",
    "Nongdam, P., Beleski, D. G., Tikendra, L., Dey, A., Varte, V., El Merzougui, S., et al. (2023). Orchid micropropagation using conventional semi-solid and temporary immersion systems: A review. Plants, 12(5), 1136. https://doi.org/10.3390/plants12051136",
    "Orvati Nia, F., Peeples, J., Murray, S. C., McFarland, A., Vann, T., Salehi, S., et al. (2026). A data-driven image extraction and analysis pipeline for plant phenotyping in controlled environments (preprint). bioRxiv. https://doi.org/10.64898/2026.02.25.707797",
    "Pastelín Solano, M. C., Salinas Ruíz, J., González Arnao, M. T., Castañeda Castro, O., Galindo Tovar, M. E., & Bello Bello, J. J. (2019). Evaluation of in vitro shoot multiplication and ISSR marker based assessment of somaclonal variants at different subcultures of vanilla (Vanilla planifolia Jacks). Physiology and Molecular Biology of Plants, 25(2), 561–567. https://doi.org/10.1007/s12298-019-00645-9",
    "Ravi, N., Gabeur, V., Hu, Y.-T., Hu, R., Ryali, C., Ma, T., Khedr, H., et al. (2024). SAM 2: Segment anything in images and videos. arXiv. https://arxiv.org/abs/2408.00714",
    "Regni, L., Calisti, S., Cesarini, A., Marconi, L., Proietti, P., Zollini, S., & Brigante, R. (2025). Micropropagation of blackberry and blueberry: Assessing the effects of subculture duration and explant density through the integration of traditional measurements and smartphone 3D imaging. Plant Cell, Tissue and Organ Culture, 163, 63. https://doi.org/10.1007/s11240-025-03267-0",
    "Thammasiri, K. (2015). Current status of orchid production in Thailand. Acta Horticulturae, 1078, 25–33. https://doi.org/10.17660/ActaHortic.2015.1078.2",
    "ศูนย์พันธุวิศวกรรมและเทคโนโลยีชีวภาพแห่งชาติ (ไบโอเทค), สวทช. (2563, 12 มิถุนายน). ไบโอเทค สวทช. พัฒนาระบบเพาะเลี้ยงพืชในอาหารเหลว เพิ่มกำลังการขยายพันธุ์ต้นกล้า. https://www.nstda.or.th/home/news_post/biotec-bioreactor/",
    "ศูนย์พันธุวิศวกรรมและเทคโนโลยีชีวภาพแห่งชาติ (ไบโอเทค), สวทช. (2565, 3 พฤษภาคม). ความสำเร็จในการขยายผลการผลิตต้นกล้าอินทผลัมในเชิงพาณิชย์ ด้วยเทคโนโลยีการเพาะเลี้ยงเนื้อเยื่อสู่เกษตรกรไทย. https://www.biotec.or.th/home/tissueculture-dates/",
]
for b in bib:
    p = para(b, first=0.508)
    p.paragraph_format.space_after = Pt(6)

# ---------- footer: เลขหน้ากลาง ไม่มี Version ----------
footer = doc.sections[0].footer
fel = footer._element
for t in list(fel.findall(qn("w:tbl"))):
    fel.remove(t)
for p in list(fel.findall(qn("w:p"))):
    fel.remove(p)
fp = footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _frun(text=None):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rF = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rF.set(qn(a), FONT)
    rPr.append(rF)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "28"); rPr.append(sz)
    r.append(rPr)
    if text is not None:
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
    return r


r1 = _frun(); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1.append(f1)
r2 = _frun(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "; r2.append(it)
r3 = _frun(); f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "separate"); r3.append(f3)
r4 = _frun("1")
r5 = _frun(); f5 = OxmlElement("w:fldChar"); f5.set(qn("w:fldCharType"), "end"); r5.append(f5)
for rr in (r1, r2, r3, r4, r5):
    fp._p.append(rr)

doc.save(OUT)
print("[OK] saved:", OUT)
d2 = Document(OUT)
print("paragraphs:", len(d2.paragraphs), "| tables:", len(d2.tables))
print("footer:", [p.text for p in d2.sections[0].footer.paragraphs if p.text.strip()])
import zipfile
z = zipfile.ZipFile(OUT)
print("media:", [n for n in z.namelist() if n.startswith("word/media")])
